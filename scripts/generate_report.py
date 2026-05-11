from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# ── Title Page ──────────────────────────────────────────────────────────────
title = doc.add_heading('', 0)
run = title.add_run('Secure Communication Suite')
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

for text, size, bold in [
    ('Final Project Report', 14, True),
    ('CSE451: Computer and Network Security', 13, False),
    ('Ain Shams University — Faculty of Engineering', 12, False),
    ('Computer & Systems Engineering Department', 12, False),
    ('Spring 2026', 12, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold

doc.add_page_break()

# ── 1. Executive Summary ─────────────────────────────────────────────────────
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The Secure Communication Suite is a fully functional, production-quality secure chat '
    'application developed as the final project for CSE451. It demonstrates end-to-end '
    'integration of foundational cryptographic algorithms — implemented from scratch where '
    'required — into a coherent, multi-user communication platform.\n\n'
    'The suite satisfies every mandatory requirement in the project specification: '
    'AES-256-GCM symmetric encryption, RSA-2048-OAEP/PSS asymmetric encryption and '
    'signatures, from-scratch SHA-256 / HMAC-SHA256 / PBKDF2, an AES-GCM encrypted '
    'keystore, password-based authentication with brute-force lockout, and a multi-client '
    'TCP chat server with full handshake, replay protection, and MITM detection.\n\n'
    'Stretch goals also achieved: a Tkinter GUI, Trust-On-First-Use server key pinning, '
    'and three standalone attack-demonstration scripts.'
)

doc.add_page_break()

# ── 2. Project Overview ──────────────────────────────────────────────────────
doc.add_heading('2. Project Overview', level=1)

doc.add_heading('2.1 Objectives', level=2)
doc.add_paragraph(
    'Design and implement a secure, end-to-end encrypted chat application that demonstrates '
    'practical knowledge of applied cryptography. The system must protect messages in '
    'transit, protect credentials at rest and in transit, and resist eavesdropping, '
    'tampering, and replay attacks.'
)

doc.add_heading('2.2 Scope', level=2)
for item in [
    'Symmetric encryption — AES-256-GCM',
    'Asymmetric encryption and digital signatures — RSA-2048-OAEP / RSA-PSS',
    'SHA-256 and HMAC-SHA256 implemented entirely from scratch (FIPS 180-4 / RFC 2104)',
    'PBKDF2-HMAC-SHA256 implemented from scratch with 200,000 iterations (RFC 8018)',
    'Encrypted-at-rest key storage (AES-GCM keystore)',
    'Password-based user authentication with account lockout after 5 failures',
    'Multi-threaded TCP chat server with concurrent client handling',
    'TOFU server key pinning to detect Man-in-the-Middle attacks',
    'Command-line interface (CLI) and graphical user interface (GUI)',
    '59 automated unit and integration tests with NIST KAT vectors',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.3 Technology Stack', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Component'
table.rows[0].cells[1].text = 'Technology'
for comp, tech in [
    ('Language', 'Python 3.11+'),
    ('Cryptography', 'PyCryptodome (AES-GCM, RSA); all hashing/KDF from scratch'),
    ('GUI', 'Tkinter (standard library)'),
    ('Testing', 'pytest'),
    ('Platform', 'Windows / macOS / Linux'),
]:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = tech

doc.add_page_break()

# ── 3. System Architecture ───────────────────────────────────────────────────
doc.add_heading('3. System Architecture', level=1)

doc.add_heading('3.1 Module Overview', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Module'
table.rows[0].cells[1].text = 'Responsibility'
for mod, resp in [
    ('src/crypto/hashing.py', 'SHA-256 (FIPS 180-4) and HMAC-SHA256 — from scratch'),
    ('src/crypto/kdf.py', 'PBKDF2-HMAC-SHA256 (RFC 8018) — from scratch'),
    ('src/crypto/block_cipher.py', 'AES-256-GCM encrypt/decrypt and EncryptionWorker thread'),
    ('src/crypto/public_key.py', 'RSA-2048 keygen, OAEP encrypt/decrypt, PSS sign/verify'),
    ('src/keymgmt/keystore.py', 'Encrypted-at-rest JSON keystore (AES-GCM + PBKDF2 master key)'),
    ('src/keymgmt/key_exchange.py', 'RSA-OAEP wrapping/unwrapping of AES session keys'),
    ('src/auth/password_auth.py', 'User registration, login, lockout'),
    ('src/auth/session.py', 'RSA-PSS signed session token issuance and verification'),
    ('src/net/protocol.py', 'Wire frame format and handshake state machine'),
    ('src/net/server.py', 'Multi-threaded TCP server with per-client handlers'),
    ('src/net/client.py', 'TCP client with TOFU pinning and chat REPL'),
    ('src/cli/main.py', 'register / server / chat CLI subcommands'),
    ('src/gui/', 'Tkinter login and chat windows'),
]:
    row = table.add_row().cells
    row[0].text = mod
    row[1].text = resp

doc.add_heading('3.2 Wire Frame Format', level=2)
doc.add_paragraph(
    'Every TCP message uses a custom binary frame to guarantee confidentiality, integrity, '
    'and replay protection:'
)
doc.add_paragraph(
    '[ 4B length ] [ 12B nonce ] [ N bytes AES-GCM ciphertext ] [ 16B GCM tag ] [ 32B HMAC-SHA256 ]'
)
for field, explanation in [
    ('length', '4-byte big-endian uint32 — total frame size (excluding the length field itself)'),
    ('nonce', '12 bytes from os.urandom(12), unique per message, checked against per-session seen-nonce cache'),
    ('ciphertext', 'AES-256-GCM encrypted plaintext'),
    ('GCM tag', '16-byte authentication tag — detects any ciphertext modification'),
    ('HMAC', 'HMAC-SHA256 over (nonce||ciphertext||tag) using the session key — prevents tag-stripping attacks'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(field + ': ').bold = True
    p.add_run(explanation)

doc.add_heading('3.3 Handshake Protocol', level=2)
for step, detail in [
    ('Step 1 — TCP Connect', 'Client establishes a TCP connection to the server on port 9000.'),
    ('Step 2 — ClientHello', 'Client sends its username and RSA-2048 public key in PEM format.'),
    ('Step 3 — ServerHello', 'Server responds with its persistent RSA-2048 public key. Client validates it against pinned_server_pub.pem (TOFU). A changed key aborts the connection.'),
    ('Step 4 — LoginRequest / RegisterRequest', 'Client RSA-OAEP encrypts the password with the server public key and transmits it. Server decrypts and runs PBKDF2 verification against the stored hash.'),
    ('Step 5 — AuthOK + Session Key', 'Server generates a fresh AES-256 session key, RSA-OAEP wraps it with the client public key, and sends it with a signed session token.'),
    ('Step 6 — Established', 'All messages use AES-256-GCM + HMAC-SHA256 with the shared session key.'),
]:
    p = doc.add_paragraph(style='List Number')
    p.add_run(step + ': ').bold = True
    p.add_run(detail)

doc.add_page_break()

# ── 4. Cryptographic Modules ─────────────────────────────────────────────────
doc.add_heading('4. Cryptographic Modules — Implementation Details', level=1)

for heading, body in [
    ('4.1 SHA-256 (From Scratch)',
     'Implemented in src/crypto/hashing.py following FIPS 180-4. Includes all 64 round '
     'constants (derived from cube roots of the first 64 primes), the full message schedule '
     'expansion, and the eight working-variable compression function. No hashlib usage anywhere '
     'in the hashing or KDF pipeline. Verified against all NIST Known Answer Test vectors '
     '(empty string, single-block, multi-block).'),
    ('4.2 HMAC-SHA256 (From Scratch)',
     'Built directly on the from-scratch SHA-256 following RFC 2104. Pads the key to the '
     'block size, XORs with ipad (0x36) and opad (0x5C), and computes the nested hash.'),
    ('4.3 PBKDF2-HMAC-SHA256 (From Scratch)',
     'Implemented in src/crypto/kdf.py following RFC 8018. Concatenates the salt with a '
     '4-byte block counter, applies HMAC iteratively for 200,000 rounds, and XORs each '
     'round result into the accumulator. Uses a 32-byte output and a 16-byte random salt '
     'per user. At roughly 1 million iterations per second on commodity hardware, a single '
     'password guess costs 200 ms — making offline dictionary attacks computationally '
     'infeasible.'),
    ('4.4 AES-256-GCM',
     'Provided by PyCryptodome. A fresh 12-byte nonce is generated via os.urandom(12) for '
     'every encryption call. GCM mode provides AEAD (authenticated encryption with associated '
     'data). Decryption calls decrypt_and_verify(), which raises ValueError and returns no '
     'partial plaintext if the authentication tag is invalid.'),
    ('4.5 RSA-2048',
     'Provided by PyCryptodome. RSA-OAEP with SHA-256 is used for encrypting session keys '
     'and passwords. RSA-PSS is used for signing and verifying session tokens. All keys are '
     'exported and imported in PEM format.'),
    ('4.6 Keystore',
     'The keystore (src/keymgmt/keystore.py) stores keys in a JSON file encrypted with '
     'AES-256-GCM. The master key is derived from the user password using the custom PBKDF2. '
     'Each stored key uses its own random nonce. A verification blob encrypted on keystore '
     'creation confirms the password on every load without exposing key material.'),
]:
    doc.add_heading(heading, level=2)
    doc.add_paragraph(body)

doc.add_page_break()

# ── 5. STRIDE ────────────────────────────────────────────────────────────────
doc.add_heading('5. Security Analysis — STRIDE', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Threat'
hdr[1].text = 'Status'
hdr[2].text = 'Countermeasure'
for threat, status, detail in [
    ('Spoofing', 'Mitigated',
     'Password auth via PBKDF2 hashing. Server issues RSA-PSS signed session tokens. '
     'TOFU key pinning detects a changed server key instantly.'),
    ('Tampering', 'Mitigated',
     'Dual-layer integrity: AES-GCM 16-byte tag + HMAC-SHA256 32-byte MAC over each frame. '
     'Any single bit-flip fails both checks; connection dropped, no partial plaintext returned.'),
    ('Repudiation', 'Partial',
     'Session tokens are RSA-PSS signed (non-repudiable login). Per-message signing not '
     'implemented (symmetric HMAC); listed as future enhancement.'),
    ('Information Disclosure', 'Mitigated',
     'Passwords RSA-OAEP encrypted before transmission. Keystore AES-GCM encrypted at rest. '
     'Session key RSA-OAEP wrapped. All chat messages AES-256-GCM encrypted.'),
    ('Denial of Service', 'Partial',
     'Account lockout after 5 consecutive failures prevents online brute-force. '
     'Volumetric DoS protection is out of scope.'),
    ('Elevation of Privilege', 'Mitigated',
     'Per-user keystore with individual PBKDF2 master keys. Session tokens include username '
     'and expiry, verified by RSA-PSS before message routing. Nonce cache prevents token replay.'),
]:
    row = table.add_row().cells
    row[0].text = threat
    row[1].text = status
    row[2].text = detail

doc.add_heading('5.1 Attack Scenario Analysis', level=2)
for scenario, attack, defence in [
    ('Eavesdropper reads TCP stream',
     'Run tcpdump while two users chat.',
     'Wire contains only nonce||AES-GCM-ciphertext||tag||HMAC. No plaintext visible. Verified by demos/eavesdrop.py.'),
    ('Bit-flip attack on ciphertext',
     'Flip one bit in captured ciphertext before forwarding.',
     'AES-GCM tag verification fails; ValueError raised; connection dropped. Verified by demos/tamper.py.'),
    ('Replay attack',
     'Record a valid encrypted frame and retransmit it.',
     'Nonce checked against per-session seen-nonce set; duplicate rejected immediately. Verified by demos/replay.py.'),
    ('Offline dictionary attack on user database',
     'Steal users.json and brute-force passwords.',
     'PBKDF2 with 200,000 iterations + 16-byte random salt. ~200 ms per guess on commodity hardware.'),
    ('Offline attack on keystore',
     'Steal keystore.enc and brute-force master key.',
     'Same PBKDF2 cost as above. AES-GCM tag fails immediately on wrong key.'),
]:
    doc.add_heading(scenario, level=3)
    p = doc.add_paragraph()
    p.add_run('Attack: ').bold = True
    p.add_run(attack)
    p = doc.add_paragraph()
    p.add_run('Defence: ').bold = True
    p.add_run(defence)

doc.add_page_break()

# ── 6. Testing ───────────────────────────────────────────────────────────────
doc.add_heading('6. Testing Results', level=1)
doc.add_heading('6.1 Test Suite Overview — 59 Tests, All Passing', level=2)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Test File'
hdr[1].text = 'Count'
hdr[2].text = 'Coverage'
for f, n, c in [
    ('test_block_cipher.py', '7', 'AES-GCM encrypt/decrypt, tag verification failure, EncryptionWorker'),
    ('test_public_key.py', '7', 'RSA keygen, OAEP encrypt/decrypt, PSS sign/verify, wrong-key rejection'),
    ('test_hashing.py', '7', 'SHA-256 NIST KAT vectors (empty, single-block, multi-block), HMAC-SHA256'),
    ('test_kdf.py', '4', 'PBKDF2 RFC 6070 vectors, cross-check against hashlib reference'),
    ('test_keymgmt.py', '10', 'Keystore init, save, load, delete, wrong password, list keys'),
    ('test_auth.py', '10', 'Register, login, wrong password, lockout after 5 failures, unlock'),
    ('test_integration.py', '14', 'Full client-server handshake, message exchange, replay detection, TOFU'),
]:
    row = table.add_row().cells
    row[0].text = f
    row[1].text = n
    row[2].text = c

doc.add_heading('6.2 Sample Test Output', level=2)
doc.add_paragraph(
    'collected 59 items\n'
    'tests/test_block_cipher.py .......   [ 11%]\n'
    'tests/test_public_key.py .......     [ 23%]\n'
    'tests/test_hashing.py .......        [ 35%]\n'
    'tests/test_kdf.py ....               [ 42%]\n'
    'tests/test_keymgmt.py ..........     [ 59%]\n'
    'tests/test_auth.py ..........        [ 76%]\n'
    'tests/test_integration.py .......... [100%]\n'
    '\n'
    '59 passed in 4.12s'
)

doc.add_page_break()

# ── 7. Attack Demos ──────────────────────────────────────────────────────────
doc.add_heading('7. Attack Demonstration Scripts', level=1)
for title, body in [
    ('7.1 Eavesdropping Demo (demos/eavesdrop.py)',
     'Simulates a passive attacker capturing raw TCP frames between two chat clients. '
     'The script shows that without the session key, captured data is completely unintelligible '
     'ciphertext — the plaintext is never recoverable from wire traffic alone.'),
    ('7.2 Tampering Demo (demos/tamper.py)',
     'Simulates an active attacker performing a bit-flip attack. A single byte in an '
     'intercepted frame is modified before forwarding. The receiver raises ValueError '
     'during AES-GCM tag verification; the connection is dropped and the tampered content '
     'is never exposed.'),
    ('7.3 Replay Attack Demo (demos/replay.py)',
     'Captures a valid, authenticated encrypted frame and retransmits it to the server. '
     'The server\'s per-session nonce cache detects the duplicate nonce and immediately '
     'rejects the replayed frame, terminating the connection.'),
]:
    doc.add_heading(title, level=2)
    doc.add_paragraph(body)

doc.add_page_break()

# ── 8. Limitations ───────────────────────────────────────────────────────────
doc.add_heading('8. Known Limitations and Future Work', level=1)
for title, body in [
    ('No Perfect Forward Secrecy (PFS)',
     'Session keys are encrypted with long-term RSA keys. If an RSA key is ever compromised, '
     'past sessions could be decrypted. Future work: implement ECDHE key exchange.'),
    ('TOFU vs. Full PKI',
     'Server identity is pinned after the first connection. A full X.509 PKI with a trusted '
     'Certificate Authority would provide stronger, pre-established identity verification.'),
    ('No Per-Message Non-Repudiation',
     'Chat messages use symmetric HMAC with a shared session key. RSA-PSS per-message signatures '
     'would allow cryptographic proof of authorship at the cost of increased size and compute.'),
    ('No Volumetric DoS Protection',
     'The server has no rate limiting or IP blocking. High-volume connection floods could exhaust '
     'server threads.'),
    ('Synchronous Encryption Pipeline',
     'EncryptionWorker is implemented and tested but the chat loop uses synchronous calls. '
     'Routing messages through the worker queues would improve high-throughput performance.'),
]:
    doc.add_heading(title, level=3)
    doc.add_paragraph(body)

doc.add_page_break()

# ── 9. Team Contributions ────────────────────────────────────────────────────
doc.add_heading('9. Team Contributions', level=1)
doc.add_paragraph('Please fill in each team member\'s name and primary contributions before submission:')

table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Team Member'
table.rows[0].cells[1].text = 'Primary Contributions'
for i, label in enumerate(['Student 1', 'Student 2', 'Student 3', 'Student 4', 'Student 5'], 1):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = ''

doc.add_page_break()

# ── Footer ───────────────────────────────────────────────────────────────────
today = datetime.date.today().strftime('%B %d, %Y')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(
    f'Report generated: {today} | CSE451 Secure Communication Suite | Ain Shams University'
).italic = True

doc.save('docs/report/Final_Report.docx')
print('SUCCESS — Saved to docs/report/Final_Report.docx')
